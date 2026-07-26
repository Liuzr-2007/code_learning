# -*- coding: utf-8 -*-
"""Word 信息提取模块

主路径：表格感知解析器，针对「十佳学风优秀团体报名表」模板，
       理解"小组成员"等类别标签与其后各行数据的主从关系。
回退路径：正则匹配"关键词：值"格式，用于非模板文档。
"""
import re
import os
from pathlib import Path

from docx import Document

from config import FIELDS, TABLE_COLUMNS
from utils import logger

# ======================== 正则预编译（回退路径用） ========================

_all_keys: list[str] = []
for _, kl in FIELDS:
    _all_keys.extend(kl)
_key_pattern_str = "|".join(re.escape(k) for k in _all_keys)
_field_patterns: dict[tuple[str, str], re.Pattern] = {}
for col_name, keywords in FIELDS:
    for key in keywords:
        pattern = re.compile(
            rf"{re.escape(key)}[：:]\s*(.+?)(?=\n|({_key_pattern_str})[：:])",
            re.S
        )
        _field_patterns[(col_name, key)] = pattern


# ======================== 表格感知解析（主路径） ========================

def _cell_text(table, row_idx: int, col_idx: int) -> str:
    """安全取表格单元格文本，越界或异常返回空串"""
    try:
        return table.rows[row_idx].cells[col_idx].text.strip()
    except (IndexError, AttributeError):
        return ""


def _row_cells(table, row_idx: int) -> list[str]:
    """取一行的所有单元格文本（已去空白）"""
    try:
        return [c.text.strip() for c in table.rows[row_idx].cells]
    except (IndexError, AttributeError):
        return []


def _last_nonempty(cells: list[str]) -> str:
    """取一行中最后一个非空单元格（用于定位手机/QQ等末列值）"""
    for cell in reversed(cells):
        if cell:
            return cell
    return ""


def _extract_table_people(doc, file_name: str) -> list[dict[str, str]] | None:
    """从「十佳学风优秀团体报名表」模板表格提取人员信息

    返回人员记录列表（组长在前，成员在后）；非此模板返回 None。
    """
    if not doc.tables:
        return None

    table = doc.tables[0]
    # 模板判定：首行首列应为"小组名称"
    if _cell_text(table, 0, 0) != "小组名称":
        return None

    rows = table.rows
    n = len(rows)

    # 1) 组名：找 col0=='小组名称' 的行，取 col1
    group_name = ""
    for i in range(n):
        if _cell_text(table, i, 0) == "小组名称":
            group_name = _cell_text(table, i, 1)
            break

    # 2) 组长：找 col0=='组长' 的行
    leader_name = ""
    leader_xuehao = ""
    leader_phone = ""
    leader_qq = ""
    for i in range(n):
        cells = _row_cells(table, i)
        if not cells or cells[0] != "组长":
            continue
        # col1 是字段标签：'姓名' 或 '学号'
        label = cells[1] if len(cells) > 1 else ""
        if label == "姓名":
            # 组长姓名在标签后一列（col2），手机在行末非空单元格
            leader_name = cells[2] if len(cells) > 2 else ""
            leader_phone = _last_nonempty(cells[2:]) if len(cells) > 2 else ""
            # 手机通常是纯数字/带+，若末列是"手机"标签则再往前找
            if leader_phone in ("手机", "联系方式"):
                leader_phone = ""
        elif label == "学号":
            leader_xuehao = cells[2] if len(cells) > 2 else ""
            leader_qq = _last_nonempty(cells[2:]) if len(cells) > 2 else ""
            if leader_qq in ("QQ", "联系方式"):
                leader_qq = ""

    # 3) 成员：先找成员表头行（col0=='小组成员' 且 col1=='姓名'），
    #    从表头行扫描 '姓名'/'学号' 出现的列索引，兼容列偏移
    member_name_col = 1
    member_xuehao_col = 4
    header_row_idx = -1
    for i in range(n):
        cells = _row_cells(table, i)
        if len(cells) >= 2 and cells[0] == "小组成员" and cells[1] == "姓名":
            header_row_idx = i
            # 找学号列：表头行中值为'学号'的列
            for ci, val in enumerate(cells):
                if val == "学号":
                    member_xuehao_col = ci
                    break
            # 姓名列默认为 col1
            member_name_col = 1
            break

    # 4) 从表头行下一行起，遍历 col0=='小组成员' 的数据行
    members: list[dict[str, str]] = []
    start = header_row_idx + 1 if header_row_idx >= 0 else 0
    for i in range(start, n):
        cells = _row_cells(table, i)
        if not cells or cells[0] != "小组成员":
            # 一旦离开"小组成员"区块即停止（后续是小组日志等其它内容）
            if members:
                break
            continue
        name = cells[member_name_col] if len(cells) > member_name_col else ""
        # 跳过表头/空行
        if not name or name == "姓名":
            continue
        xuehao = cells[member_xuehao_col] if len(cells) > member_xuehao_col else ""
        # 去重：组长也列在成员表中则跳过
        if leader_name and name == leader_name:
            continue
        members.append({
            "文档名称": file_name,
            "小组名称": group_name,
            "角色": "小组成员",
            "姓名": name,
            "学号": xuehao,
            "手机": "",
            "QQ": "",
        })

    # 5) 组装结果：组长在前，成员在后
    people: list[dict[str, str]] = []
    if leader_name:
        people.append({
            "文档名称": file_name,
            "小组名称": group_name,
            "角色": "组长",
            "姓名": leader_name,
            "学号": leader_xuehao,
            "手机": leader_phone,
            "QQ": leader_qq,
        })
    people.extend(members)

    # 模板识别但无人（极端情况）→ 返回空列表（非 None），表示已处理
    return people


# ======================== 正则回退提取（非模板文档） ========================

def _extract_by_regex(content: str, file_name: str) -> list[dict[str, str]]:
    """正则匹配"关键词：值"格式，映射到 TABLE_COLUMNS 兼容字典"""
    field_matches: dict[str, list[str]] = {}
    for col_name, keywords in FIELDS:
        all_vals: list[str] = []
        for key in keywords:
            pattern = _field_patterns[(col_name, key)]
            for match in pattern.finditer(content):
                all_vals.append(match.group(1).strip())
        if all_vals:
            seen: set[str] = set()
            unique_vals: list[str] = []
            for v in all_vals:
                if v not in seen:
                    seen.add(v)
                    unique_vals.append(v)
            field_matches[col_name] = unique_vals

    names = field_matches.get("姓名", [])
    if not names:
        return [{
            "文档名称": file_name, "小组名称": "", "角色": "",
            "姓名": "", "学号": "", "手机": "", "QQ": "",
        }]

    phones = field_matches.get("手机号", [])
    results: list[dict[str, str]] = []
    for i, name in enumerate(names):
        results.append({
            "文档名称": file_name,
            "小组名称": "",
            "角色": "",
            "姓名": name,
            "学号": "",
            "手机": phones[i] if i < len(phones) else "",
            "QQ": "",
        })
    return results


# ======================== 对外接口 ========================

def extract_word_info(file_path: str) -> list[dict[str, str]]:
    """读取单个 docx 提取人员信息，每人返回一条记录（TABLE_COLUMNS schema）

    优先用表格感知解析；非模板文档回退到正则匹配。
    """
    file_name = os.path.basename(file_path)

    try:
        doc = Document(file_path)
    except Exception as e:
        logger.warning(f"文档 {file_name} 损坏/无法读取，跳过：{e}")
        return [{
            "文档名称": file_name, "小组名称": "", "角色": "",
            "姓名": "", "学号": "", "手机": "", "QQ": "",
        }]

    # 主路径：表格感知解析
    people = _extract_table_people(doc, file_name)
    if people is not None:
        return people

    # 回退路径：正则匹配"关键词：值"
    content = ""
    for p in doc.paragraphs:
        content += p.text + "\n"
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                content += cell.text + "\n"
    return _extract_by_regex(content, file_name)


def scan_word_folder(folder_path: str) -> list[dict[str, str]]:
    """遍历文件夹收集全部 docx，返回所有人员信息（每人一行）"""
    word_files: list[str] = []
    for root, _, files in os.walk(folder_path):
        current_root = Path(root)
        for f in files:
            if f.lower().endswith(".docx"):
                word_files.append(str(current_root / f))

    all_data: list[dict[str, str]] = []
    for file in word_files:
        print(f"正在读取：{os.path.basename(file)}")
        data_list = extract_word_info(file)
        all_data.extend(data_list)
    return all_data
