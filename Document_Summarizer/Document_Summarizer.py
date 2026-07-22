# -*- coding: utf-8 -*-
import os
import re
import sys
import ctypes
import logging
import shutil
import zipfile
import tempfile
import datetime
import subprocess
from pathlib import Path
from collections import defaultdict
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from PIL import Image
from docx import Document
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

# 进度条：tqdm 为可选依赖，未安装时降级为普通迭代，不影响功能
try:
    from tqdm import tqdm
except ImportError:
    def tqdm(iterable=None, **kwargs):
        return iterable

# 日志配置
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(message)s",
    handlers=[logging.StreamHandler()]
)
logger = logging.getLogger(__name__)

# Windows 文件属性标志常量
FILE_ATTRIBUTE_READONLY = 0x0001  # 只读文件
FILE_ATTRIBUTE_HIDDEN = 0x0002    # 隐藏文件
FILE_ATTRIBUTE_SYSTEM = 0x0004    # 系统文件
FILE_ATTRIBUTE_DIRECTORY = 0x0010 # 目录/文件夹
FILE_ATTRIBUTE_ARCHIVE = 0x0020   # 存档文件

# 平台判断：非 Windows 时不绑定 kernel32，降级为“点开头文件即隐藏”的判断
IS_WINDOWS = sys.platform == "win32"
if IS_WINDOWS:
    kernel32 = ctypes.windll.kernel32
    kernel32.GetFileAttributesW.argtypes = [ctypes.c_wchar_p]
    kernel32.GetFileAttributesW.restype = ctypes.c_uint32
else:
    kernel32 = None


# ======================== 公共工具函数 ========================

def clean_path_input(raw: str) -> Path:
    """统一清理用户输入路径：去引号、去空格、展开环境变量"""
    cleaned = raw.strip().strip('"').strip("'")
    expanded = os.path.expandvars(cleaned)
    return Path(expanded)


def display_width(text: str) -> int:
    """计算字符串显示宽度，中文字符算2，英文算1"""
    return sum(2 if ord(c) > 127 else 1 for c in text)


def get_file_attributes(file_path: Path) -> int:
    """获取Windows文件属性标志，失败或非Windows返回-1"""
    if kernel32 is None:
        return -1
    try:
        return kernel32.GetFileAttributesW(str(file_path))
    except OSError:
        return -1


def is_hidden_file(file_path: Path) -> bool:
    """判断是否拥有隐藏属性"""
    attrs = get_file_attributes(file_path)
    if attrs == -1:
        return False
    return bool(attrs & FILE_ATTRIBUTE_HIDDEN)


def is_system_file(file_path: Path) -> bool:
    """判断是否为系统文件"""
    attrs = get_file_attributes(file_path)
    if attrs == -1:
        return False
    return bool(attrs & FILE_ATTRIBUTE_SYSTEM)


def skip_hidden_system(path: Path) -> bool:
    """综合过滤规则：隐藏/系统文件/点开头文件 返回True=跳过扫描"""
    return is_hidden_file(path) or is_system_file(path) or path.name.startswith(".")


def format_size(byte_num: int) -> str:
    """字节自动格式化 GB/MB/KB/B"""
    if byte_num >= 1024 ** 3:
        return f"{byte_num / (1024 ** 3):.2f} GB"
    elif byte_num >= 1024 ** 2:
        return f"{byte_num / (1024 ** 2):.2f} MB"
    elif byte_num >= 1024:
        return f"{byte_num / 1024:.2f} KB"
    else:
        return f"{byte_num} B"


def get_file_time_str(stat_info, time_type: str = "create") -> str:
    """统一提取文件时间"""
    try:
        if time_type == "create":
            ts = stat_info.st_ctime
        else:
            ts = stat_info.st_mtime
        dt = datetime.datetime.fromtimestamp(ts)
        return dt.strftime("%Y-%m-%d %H:%M:%S")
    except (OSError, ValueError) as e:
        return f"读取失败: {str(e)}"


def get_unique_path(base_path: Path) -> Path:
    """如果文件已存在，自动编号 (1), (2), ..."""
    if not base_path.exists():
        return base_path
    stem = base_path.stem
    suffix = base_path.suffix
    parent = base_path.parent
    counter = 1
    while True:
        new_path = parent / f"{stem}({counter}){suffix}"
        if not new_path.exists():
            return new_path
        counter += 1


# ======================== 功能一：文件夹扫描 ========================

def collect_folder_info(folder_input: str) -> None:
    """遍历文件夹生成扫描统计txt报告"""
    target = clean_path_input(folder_input)

    if not target.is_dir():
        logger.error(f"路径「{target}」不是有效文件夹！")
        return

    while True:
        choose_hidden = input("是否一并查询隐藏/系统文件？输入 y 包含，n 过滤：").strip().lower()
        if choose_hidden in ("y", "n"):
            break
        print("无效输入，请输入 y 或 n")

    # 记录统一扫描时间（修复双重 datetime.now() 问题）
    scan_time = datetime.datetime.now()
    time_suffix = scan_time.strftime("%Y%m%d_%H%M%S")

    print(f"\n==========开始扫描目录: {target} ==========")
    all_files: list[dict] = []
    suffix_counter: dict[str, int] = defaultdict(int)
    total_byte = 0
    dir_count = 0
    read_fail_count = 0

    for root, dirs, files in os.walk(target):
        current_root = Path(root)
        print(f"正在扫描目录: {current_root}")

        if choose_hidden == "n":
            dirs[:] = [d for d in dirs if not skip_hidden_system(current_root / d)]
        dir_count += len(dirs)

        for filename in files:
            file_full = current_root / filename
            if choose_hidden == "n" and skip_hidden_system(file_full):
                continue

            # 修复：初始化 stat_info 防止未定义崩溃
            stat_info = None
            try:
                stat_info = file_full.stat()
                file_size = stat_info.st_size
                total_byte += file_size
                err_msg = None
            except (PermissionError, FileNotFoundError, OSError) as e:
                file_size = 0
                err_msg = str(e)
                read_fail_count += 1

            # 修复：用 stat_info is not None 判断，而非 err_msg
            create_t = get_file_time_str(stat_info, "create") if stat_info is not None else err_msg
            mod_t = get_file_time_str(stat_info, "modify") if stat_info is not None else err_msg
            suffix = file_full.suffix if file_full.suffix else "无后缀"
            suffix_counter[suffix] += 1

            file_info = {
                "name": filename,
                "path": str(file_full),
                "size_byte": file_size,
                "size_str": format_size(file_size) if stat_info is not None else err_msg,
                "create_time": create_t,
                "modify_time": mod_t
            }
            all_files.append(file_info)
            print(f"【文件】{file_info['name']} | {file_info['size_str']}")
            print(f"完整路径: {file_info['path']}")
            print(f"创建: {create_t} | 修改: {mod_t}")
            print("-" * 60)

    # 报告保存到扫描目标文件夹内
    report_name = f"文件夹扫描报告_{time_suffix}.txt"
    report_path = get_unique_path(target / report_name)

    with open(report_path, "w", encoding="utf-8") as f:
        scan_mode = "（包含隐藏/系统文件）" if choose_hidden == "y" else "（已过滤隐藏/系统文件）"
        f.write(f"===== 文件夹扫描汇总报告 {scan_mode} =====\n")
        f.write(f"扫描根目录：{target}\n")
        f.write(f"扫描时间：{scan_time.strftime('%Y-%m-%d %H:%M:%S')}\n")
        f.write(f"子文件夹总数：{dir_count} 个\n")
        f.write(f"有效文件总数：{len(all_files)} 个\n")
        f.write(f"读取失败文件：{read_fail_count} 个\n")
        f.write(f"总占用空间：{format_size(total_byte)}\n\n")

        f.write("===== 文件后缀分类统计 =====\n")
        sorted_suffix = sorted(suffix_counter.items(), key=lambda x: x[1], reverse=True)
        for ext, count in sorted_suffix:
            f.write(f"{ext:10s}：{count:4d} 个\n")

        f.write("\n===== 全部文件明细列表 =====\n")
        for info in all_files:
            f.write("-" * 60 + "\n")
            f.write(f"文件名：{info['name']}\n")
            f.write(f"完整路径：{info['path']}\n")
            f.write(f"文件大小：{info['size_str']}\n")
            f.write(f"创建时间：{info['create_time']}\n")
            f.write(f"修改时间：{info['modify_time']}\n")

    print(f"\n==================== 扫描完成 ====================")
    print(f"根目录子文件夹：{dir_count} 个")
    print(f"扫描到文件总数：{len(all_files)} 个")
    print(f"读取失败文件：{read_fail_count} 个")
    print(f"目录总占用空间：{format_size(total_byte)}")
    print(f"完整报告已保存至：{report_path}")


# ======================== 功能二：图片合并PDF ========================

def folder_images_to_pdf(folder_path: str) -> None:
    """文件夹图片批量生成A4 PDF，递归子目录"""
    target = clean_path_input(folder_path)
    if not target.is_dir():
        logger.error("路径不是有效文件夹！")
        return

    output_pdf = str(get_unique_path(target / "批量照片汇总.pdf"))
    img_suffix = (".jpg", ".jpeg", ".png", ".bmp", ".tiff")
    img_path_list: list[str] = []

    # 递归遍历子目录收集图片
    for root, _, files in os.walk(target):
        current_root = Path(root)
        for filename in files:
            if filename.lower().endswith(img_suffix):
                full_file_path = current_root / filename
                if full_file_path.is_file():
                    img_path_list.append(str(full_file_path))

    if len(img_path_list) == 0:
        logger.warning("未检测到任何图片文件，PDF生成失败")
        return

    # 按文件名排序，确保输出顺序稳定
    img_path_list.sort(key=lambda p: Path(p).name.lower())

    pdf = canvas.Canvas(output_pdf, pagesize=A4)
    page_width, page_height = A4

    for img_path in tqdm(img_path_list, desc="合并PDF"):
        try:
            draw_img_path = img_path
            temp_file = None

            with Image.open(img_path) as img:
                img_width, img_height = img.size

                # 修复：RGBA PNG 转 RGB 兼容 reportlab
                if img.mode == "RGBA":
                    rgb_img = img.convert("RGB")
                    temp_fd, temp_file = tempfile.mkstemp(suffix=".jpg")
                    os.close(temp_fd)
                    rgb_img.save(temp_file, "JPEG")
                    draw_img_path = temp_file

                scale_rate = min(page_height / img_height, page_width / img_width)
                draw_w = img_width * scale_rate
                draw_h = img_height * scale_rate
                X = (page_width - draw_w) / 2
                Y = (page_height - draw_h) / 2

            pdf.drawImage(draw_img_path, X, Y, width=draw_w, height=draw_h)
            pdf.showPage()

            # 清理临时文件
            if temp_file and os.path.exists(temp_file):
                os.remove(temp_file)

        except (OSError, ValueError, RuntimeError) as e:
            logger.error(f"图片 {img_path} 处理失败: {str(e)}")
            # 清理临时文件
            if 'temp_file' in locals() and temp_file and os.path.exists(temp_file):
                os.remove(temp_file)

    pdf.save()
    print(f"PDF生成成功，路径：{output_pdf}")


# ======================== 功能三：Word提取信息 ========================

# Word提取配置区
FIELDS: list[tuple[str, list[str]]] = [
    ("姓名", ["姓名", "姓 名"]),
    ("性别", ["性别", "性 别"]),
    ("年龄", ["年龄"]),
    ("身份证号", ["身份证", "身份证号", "证件号"]),
    ("手机号", ["电话", "手机号", "联系电话"]),
    ("住址", ["住址", "地址", "家庭住址"]),
    ("学历", ["学历"]),
    ("工作单位", ["单位", "工作单位"])
]

# 预编译正则：避免循环内重复编译
_all_keys: list[str] = []
for _, kl in FIELDS:
    _all_keys.extend(kl)
_key_pattern_str = "|".join(re.escape(k) for k in _all_keys)
_field_patterns: dict[str, re.Pattern] = {}
for col_name, keywords in FIELDS:
    for key in keywords:
        pattern = re.compile(
            rf"{re.escape(key)}[：:]\s*(.+?)(?=\n|({_key_pattern_str})[：:])",
            re.S
        )
        _field_patterns[(col_name, key)] = pattern


def convert_doc_to_docx(doc_path: str) -> str | None:
    """调用 LibreOffice 将旧版 .doc 转换为 .docx，返回临时 docx 路径；失败返回 None"""
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if not soffice:
        logger.warning(
            "未检测到 LibreOffice(soffice)，无法读取 .doc 文件，"
            "请安装 LibreOffice 或将文档另存为 .docx 后重试"
        )
        return None
    out_dir = tempfile.mkdtemp(prefix="doc_conv_")
    try:
        subprocess.run(
            [soffice, "--headless", "--convert-to", "docx",
             "--outdir", out_dir, doc_path],
            check=True, timeout=120,
            stdout=subprocess.DEVNULL, stderr=subprocess.PIPE,
        )
    except (subprocess.CalledProcessError, subprocess.TimeoutExpired, OSError) as e:
        logger.warning(f"LibreOffice 转换「{Path(doc_path).name}」失败：{e}")
        shutil.rmtree(out_dir, ignore_errors=True)
        return None
    out_docx = Path(out_dir) / (Path(doc_path).stem + ".docx")
    if not out_docx.exists():
        shutil.rmtree(out_dir, ignore_errors=True)
        return None
    return str(out_docx)


def extract_word_info(file_path: str) -> list[dict[str, str]]:
    """读取单个 docx/doc 提取人员信息，支持多人，每人返回一条记录"""
    file_name = Path(file_path).name
    empty_data = {col: "" for col, _ in FIELDS}

    actual_path = file_path
    temp_dir: str | None = None
    # 旧版 .doc 先用 LibreOffice 转换为 .docx
    if file_path.lower().endswith(".doc"):
        converted = convert_doc_to_docx(file_path)
        if converted is None:
            return [{"文档名称": file_name, **empty_data}]
        actual_path = converted
        temp_dir = str(Path(converted).parent)

    try:
        # 收窄异常类型并记录 traceback，避免静默吞掉真正的问题
        try:
            doc = Document(actual_path)
        except (zipfile.BadZipFile, KeyError, IndexError, OSError, ValueError) as e:
            logger.warning(f"文档「{file_name}」损坏/无法读取，跳过：{e}", exc_info=True)
            return [{"文档名称": file_name, **empty_data}]

        content = ""
        # 读取段落
        for p in doc.paragraphs:
            content += p.text + "\n"
        # 读取表格
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    content += cell.text + "\n"

        # 对每个字段用 finditer 找出所有匹配值
        field_matches: dict[str, list[str]] = {}
        for col_name, keywords in FIELDS:
            all_vals: list[str] = []
            for key in keywords:
                pattern = _field_patterns[(col_name, key)]
                for match in pattern.finditer(content):
                    all_vals.append(match.group(1).strip())
            if all_vals:
                # 去重（保持顺序）
                seen: set[str] = set()
                unique_vals: list[str] = []
                for v in all_vals:
                    if v not in seen:
                        seen.add(v)
                        unique_vals.append(v)
                field_matches[col_name] = unique_vals

        # 确定人数（取各字段最大匹配数）
        max_persons = max((len(v) for v in field_matches.values()), default=0)
        if max_persons == 0:
            return [{"文档名称": file_name, **empty_data}]

        # 每人一条记录
        results: list[dict[str, str]] = []
        for i in range(max_persons):
            person: dict[str, str] = {"文档名称": file_name}
            for col_name, _ in FIELDS:
                vals = field_matches.get(col_name, [])
                person[col_name] = vals[i] if i < len(vals) else ""
            results.append(person)
        return results
    finally:
        # 清理 .doc 转换产生的临时目录
        if temp_dir:
            shutil.rmtree(temp_dir, ignore_errors=True)


def scan_word_folder(folder_path: str) -> list[dict[str, str]]:
    """遍历文件夹收集全部 docx/doc，返回所有人员信息（每人一行）"""
    word_files: list[str] = []
    for root, _, files in os.walk(folder_path):
        current_root = Path(root)
        for f in files:
            low = f.lower()
            if low.endswith(".docx") or low.endswith(".doc"):
                word_files.append(str(current_root / f))

    all_data: list[dict[str, str]] = []
    for file in tqdm(word_files, desc="读取Word文档"):
        print(f"正在读取：{Path(file).name}")
        data_list = extract_word_info(file)
        all_data.extend(data_list)
    return all_data


def export_excel(save_path: str, data_list: list[dict[str, str]]) -> None:
    """导出汇总Excel，中文列宽自适应"""
    wb = Workbook()
    ws = wb.active
    ws.title = "人员信息汇总"

    headers = ["文档名称"] + [item[0] for item in FIELDS]
    ws.append(headers)

    # 表头样式
    header_fill = PatternFill(start_color="CCE5FF", end_color="CCE5FF", fill_type="solid")
    head_font = Font(bold=True)
    thin_border = Border(
        left=Side(style="thin"),
        right=Side(style="thin"),
        top=Side(style="thin"),
        bottom=Side(style="thin")
    )

    for col in range(1, len(headers) + 1):
        cell = ws.cell(row=1, column=col)
        cell.fill = header_fill
        cell.font = head_font
        cell.alignment = Alignment(horizontal="center", vertical="center")
        cell.border = thin_border

    # 写入数据行
    for row_data in data_list:
        row = [row_data["文档名称"]] + [row_data[col] for col, _ in FIELDS]
        ws.append(row)

    # 为数据行添加边框
    for row in ws.iter_rows(min_row=2, max_row=ws.max_row, min_col=1, max_col=len(headers)):
        for cell in row:
            cell.border = thin_border

    # 自动列宽（修复：使用 display_width 适应中文字符）
    for col_cells in ws.columns:
        max_len = 0
        col_letter = col_cells[0].column_letter
        for cell in col_cells:
            if cell.value:
                max_len = max(max_len, display_width(str(cell.value)))
        ws.column_dimensions[col_letter].width = max_len + 4

    wb.save(save_path)
    print(f"\nExcel文件已保存至：{save_path}")


def word_to_excel_main(folder_input: str) -> None:
    """Word转Excel主逻辑"""
    target = clean_path_input(folder_input)
    if not target.is_dir():
        logger.error("输入路径不是有效文件夹！")
        return
    # Excel保存到目标文件夹根目录，自动编号避免覆盖
    excel_save = str(get_unique_path(target / "人员信息汇总.xlsx"))
    data = scan_word_folder(str(target))
    if len(data) == 0:
        logger.warning("文件夹内未找到任何docx文档！")
        return
    export_excel(excel_save, data)


# ======================== 主菜单 ========================

def print_menu() -> None:
    """打印主菜单"""
    print("\n==================== 文件工具主菜单 ====================")
    print("[CFI]        文件夹深度扫描统计（生成txt报告）")
    print("[PDF]        文件夹图片批量合并PDF")
    print("[WORD2EXCEL] 读取文件夹所有Word提取人员信息生成Excel")
    print("[q]          退出程序")
    print("===========================================================\n")


if __name__ == "__main__":
    print_menu()
    argv_used = False  # 标记拖拽路径是否已使用

    while True:
        choice = input("请输入功能指令：").strip().upper()

        if choice == "Q":
            print("程序正常退出")
            sys.exit(0)
        elif choice == "CFI":
            if not argv_used and len(sys.argv) >= 2:
                folder_input = sys.argv[1]
                argv_used = True
                print(f"检测到拖拽路径：{folder_input}")
            else:
                folder_input = input("请输入目标文件夹路径（可直接拖拽文件夹到窗口）：")
            collect_folder_info(folder_input)
            input("\n按回车键返回功能菜单...")
            print_menu()
        elif choice == "PDF":
            if not argv_used and len(sys.argv) >= 2:
                folder_input = sys.argv[1]
                argv_used = True
                print(f"检测到拖拽路径：{folder_input}")
            else:
                folder_input = input("请输入图片文件夹路径：")
            folder_images_to_pdf(folder_input)
            input("\n按回车键返回功能菜单...")
            print_menu()
        elif choice == "WORD2EXCEL":
            if not argv_used and len(sys.argv) >= 2:
                folder_input = sys.argv[1]
                argv_used = True
                print(f"检测到拖拽路径：{folder_input}")
            else:
                folder_input = input("请输入存放Word文档的文件夹路径：")
            word_to_excel_main(folder_input)
            input("\n按回车键返回功能菜单...")
            print_menu()
        else:
            print("无效指令，请重新输入！")
            input("按回车键返回菜单")
            print_menu()
